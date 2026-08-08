"""Render the implementation whitepaper Markdown to a styled Word document.

No pandoc on this machine, so this does the conversion directly with python-docx:
headings, real Word tables, inline bold/italic/code, lists, blockquotes, code
blocks, a title page and an auto-populating table of contents.

Mermaid blocks cannot be rendered to images without a browser toolchain, so they
are emitted as labelled figures with the diagram source in monospace.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\Kuziwa\Desktop\Treasury\Treasury-ALM-Risk-Platform-Implementation-Whitepaper.md"
OUT = r"C:\Users\Kuziwa\Desktop\Treasury\Treasury-ALM-Risk-Platform-Implementation-Whitepaper.docx"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)      # deep navy
ACCENT2 = RGBColor(0x2E, 0x6B, 0x8A)     # steel blue
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
HDR_SHADE = "1F3A5F"
ALT_SHADE = "F2F5F8"


# --------------------------------------------------------------------------- setup

def setup(doc):
    for s in doc.sections:
        s.page_height, s.page_width = Cm(29.7), Cm(21.0)
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")

    specs = [
        ("Heading 1", 20, True, ACCENT, 20, 8),
        ("Heading 2", 15, True, ACCENT, 15, 5),
        ("Heading 3", 12.5, True, ACCENT2, 12, 4),
        ("Heading 4", 11, True, ACCENT2, 10, 3),
    ]
    for name, size, bold, colour, before, after in specs:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def shade(cell, hexcolour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolour)
    cell._tc.get_or_add_tcPr().append(el)


def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t"); txt.text = "Right-click and choose Update Field to build the contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for e in (fld, instr, sep, txt, end):
        r._r.append(e)


# --------------------------------------------------------------- inline formatting

TOKEN = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|(?<![\*\w])\*[^\*]+?\*(?!\*))")


def emit_inline(par, text, base_bold=False):
    for piece in TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8A, 0x2B, 0x2B)
        elif piece.startswith("*") and piece.endswith("*"):
            r = par.add_run(piece[1:-1]); r.italic = True
        else:
            r = par.add_run(piece)
        if base_bold:
            r.bold = True


# ------------------------------------------------------------------------ elements

def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    for i, cell_text in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        emit_inline(p, cell_text, base_bold=True)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
        shade(c, HDR_SHADE)

    for n, row in enumerate(body):
        cells = t.add_row().cells
        for i in range(len(header)):
            val = row[i] if i < len(row) else ""
            c = cells[i]
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            emit_inline(p, val)
            for run in p.runs:
                run.font.size = Pt(9.5)
            if n % 2 == 1:
                shade(c, ALT_SHADE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


IMG_DIR = r"C:\Users\Kuziwa\Desktop\Treasury\diagrams"

# Rendered diagrams, in the order the mermaid blocks appear in the Markdown.
#            file,                              width cm, caption
FIGURES = [
    ("01-domain-map.png", 16.0,
     "Seventeen bounded contexts in six layers. Solid = data flow; dashed = rule, state and query "
     "edges. Principal edges shown; the full set is enumerated in Part II §2"),
    ("02-eod-dag.png", 11.5,
     "The end-of-day pipeline as a directed acyclic graph. A stage blocks only its descendants"),
    ("03-tier-a-critical-path.png", 16.0,
     "The tier A critical path — what must exist by 07:00, and what it deliberately avoids"),
    ("04-phase0-dependency-graph.png", 15.5,
     "Phase 0 ticket dependencies across six waves"),
]

_fig_counter = {"n": 0}


def add_figure(doc, index):
    """Embed a pre-rendered diagram with a numbered caption."""
    import os
    if index >= len(FIGURES):
        return False
    fname, width_cm, caption = FIGURES[index]
    path = os.path.join(IMG_DIR, fname)
    if not os.path.exists(path):
        return False

    _fig_counter["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    # Printable width at 2.2cm margins on A4 is 16.6cm; each figure is sized
    # so its rendered height also fits the page.
    p.add_run().add_picture(path, width=Cm(width_cm))

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run(f"Figure {_fig_counter['n']} — {caption}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    return True


def add_code(doc, lines, lang):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_sep(line):
    return bool(re.fullmatch(r"\|?[\s:\-\|]+\|?", line.strip())) and "-" in line


# --------------------------------------------------------------------------- build

def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    setup(doc)

    # ---- title page
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Treasury, ALM & Risk Platform")
    r.font.size = Pt(30); r.bold = True; r.font.color.rgb = ACCENT

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Implementation Whitepaper")
    r.font.size = Pt(17); r.font.color.rgb = ACCENT2

    doc.add_paragraph()
    for line in ("Ready for implementation · Phase 0 executable",
                 "Scope basis: Tier 1 Bank Treasury — Instrument Universe & Granular Balance Sheet Taxonomy",
                 "30 artifacts · 17 bounded contexts · one adversarial review · two independent taxonomy validations"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(10); r.font.color.rgb = MUTED

    doc.add_page_break()
    h = doc.add_paragraph(); h.style = doc.styles["Heading 1"]
    h.add_run("Contents")
    add_toc(doc)
    doc.add_page_break()

    i = 0
    started = False          # skip the markdown title block
    pending_break = False
    mermaid_seen = [0]       # index into FIGURES, in document order
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            buf, i = [], i + 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            if lang == "mermaid":
                # Embed the pre-rendered PNG; fall back to source if missing.
                if not add_figure(doc, mermaid_seen[0]):
                    add_code(doc, buf, lang)
                mermaid_seen[0] += 1
            else:
                add_code(doc, buf, lang)
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            rows = [split_row(stripped)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            add_table(doc, rows)
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2)
            # Everything before "How to use this document" is title-page material
            # and is already rendered there; skip it rather than repeat it.
            if not started:
                if text.strip().lower().startswith("how to use"):
                    started = True
                else:
                    i += 1
                    continue
            if level == 1:
                doc.add_page_break()
            elif pending_break:
                doc.add_page_break()
            pending_break = False
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {level}"]
            emit_inline(p, re.sub(r"^#+\s*", "", text))
            i += 1
            continue

        if not started:
            i += 1
            continue

        # horizontal rule -> page break before next heading
        if stripped in ("---", "***", "___"):
            pending_break = True
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            body = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            emit_inline(p, body)
            for r in p.runs:
                r.italic = True
                if r.font.color.rgb is None:
                    r.font.color.rgb = ACCENT2
            i += 1
            continue

        # lists
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            depth = min(len(m.group(1)) // 2, 2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + 0.5 * depth)
            p.paragraph_format.space_after = Pt(3)
            emit_inline(p, m.group(2))
            i += 1
            continue

        m = re.match(r"^(\s*)\d+[.)]\s+(.*)$", line)
        if m:
            depth = min(len(m.group(1)) // 2, 2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.6 + 0.5 * depth)
            p.paragraph_format.space_after = Pt(3)
            emit_inline(p, m.group(2))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        emit_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("Written:", OUT)


if __name__ == "__main__":
    main()
